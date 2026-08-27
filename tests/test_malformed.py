from pathlib import Path
import pytest
from bs4 import BeautifulSoup

from okfguard.api import sanitize
from docx import Document
from docx.oxml import parse_xml

def test_html_important(tmp_path):
    """Test what happens when CSS uses !important."""
    html = '''
    <html><head><style>.hidden { display: none !important; }</style></head>
    <body><div class="hidden">Legitimate hidden note.</div></body></html>
    '''
    f = tmp_path / "important.html"
    f.write_text(html)
    
    res = sanitize(str(f))
    # !important is now properly stripped, so the hiding CSS is evaluated correctly and quarantined.
    assert res.action == "quarantine", "Expected it to detect display: none despite the !important override."

def test_html_no_head_style(tmp_path):
    """Test what happens with no head/style tags."""
    html = '<div style="display: none;">Legitimate hidden note.</div>'
    f = tmp_path / "no_head.html"
    f.write_text(html)
    
    res = sanitize(str(f))
    print(res.flags)
    assert res.action == "quarantine", "Expected it to still detect inline styles without head."

def test_docx_no_shd(tmp_path):
    """Test DOCX XML parsing fallback when w:tcPr exists but w:shd is missing."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = "Normal text"
    
    # Manually add w:tcPr but no w:shd
    tcPr = parse_xml('<w:tcPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    cell._tc.append(tcPr)
    
    f = tmp_path / "no_shd.docx"
    doc.save(f)
    
    # If the XML parser throws an exception, it exits with 3 (raises in api).
    # If it falls back gracefully, it returns clean text.
    res = sanitize(str(f))
    assert res.action == "pass", "Expected graceful fallback to pass."

def test_docx_bad_font_rgb(tmp_path):
    """Test DOCX font rgb extraction ValueError fallback."""
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Normal text")
    
    # python-docx font.color.rgb is a shared descriptor. Let's patch it for this run to raise ValueError.
    # Actually, we can just mock it or test the internal function directly.
    from okfguard.adapters.docx import _get_run_colour_rgb
    
    class BadRun:
        @property
        def font(self):
            class BadFont:
                @property
                def color(self):
                    class BadColor:
                        @property
                        def rgb(self):
                            return "invalid_hex"
                    return BadColor()
            return BadFont()
            
    res = _get_run_colour_rgb(BadRun())
    assert res is None, "Expected graceful fallback to None on ValueError."

