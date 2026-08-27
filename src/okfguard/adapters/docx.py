"""Adapter for Word documents (.docx) — hidden-text and shading detection.

Word documents have an explicit hidden-text mechanism (``font.hidden``)
that is the highest-confidence hidden-content signal across any adapter
in this project, since it is an explicit document property — not
inferred from rendering.

Three independent checks are performed:

1. **``run.font.hidden``** — ``python-docx`` exposes this directly as a
   boolean.  When ``True``, the run is genuinely hidden text per Word's
   own formatting model.  Confidence: 0.9 (explicit).

2. **White-text-on-white-background** — the run's font colour is
   compared against the effective background (paragraph/cell shading
   fill, defaulting to white).  Confidence: 0.75 (inferred).

3. **Paragraph / cell shading matching text colour** — retrieved
   directly from the underlying XML via ``docx.oxml`` rather than
   relying on ``python-docx``'s high-level shading API, which can
   return ``None`` even when shading is genuinely set (the high-level
   API's ``None`` is ambiguous between "no shading" and "shading exists
   but API failed to read it" — going to the XML eliminates that
   ambiguity).

Checks 1, 2, and 3 run independently side-by-side — the shading/colour
check does not substitute for or replace the ``font.hidden`` check.

**Requires:** ``python-docx`` — install via
``pip install okf-guard[docx]``.
"""

from __future__ import annotations

import io
import os

from okfguard.adapters.base import SourceAdapter, _utc_now_iso
from okfguard.core.models import ExtractedContent

# WordprocessingML XML namespace.
_WML_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Colour tolerance per channel for the white-on-white check.
_COLOUR_TOLERANCE = 10


# ---------------------------------------------------------------------------
# Colour helpers
# ---------------------------------------------------------------------------

def _parse_hex_colour(hex_str: str | None) -> tuple[int, int, int] | None:
    """Parse a 6-digit hex colour string (e.g. ``'FFFFFF'``) to RGB.

    Returns ``None`` for unparseable or absent values.
    """
    if not hex_str:
        return None
    hex_str = hex_str.strip().lstrip("#")
    if len(hex_str) != 6:
        return None
    try:
        return (
            int(hex_str[0:2], 16),
            int(hex_str[2:4], 16),
            int(hex_str[4:6], 16),
        )
    except ValueError:
        return None


def _colours_close(
    c1: tuple[int, int, int],
    c2: tuple[int, int, int],
) -> bool:
    """Return ``True`` if two RGB colours are within tolerance."""
    return all(abs(a - b) <= _COLOUR_TOLERANCE for a, b in zip(c1, c2))


def _get_run_colour_rgb(run: object) -> tuple[int, int, int] | None:
    """Extract an explicit RGB colour from a python-docx Run.

    Returns ``None`` if the colour is theme-based, inherited, or not
    set — only explicit RGB values are returned, since theme colours
    cannot be confidently resolved without the full theme XML.
    """
    try:
        # python-docx objects are not fully typed.
        font = run.font  # type: ignore[attr-defined]
        if font.color is None or font.color.rgb is None:
            return None
        rgb = font.color.rgb
        val = int(rgb)
        return ((val >> 16) & 0xFF, (val >> 8) & 0xFF, val & 0xFF)
    except (AttributeError, TypeError, ValueError):
        return None


# ---------------------------------------------------------------------------
# XML-based shading extraction
# ---------------------------------------------------------------------------

def _get_shading_fill_from_xml(element: object) -> str | None:
    """Extract the ``w:fill`` colour from a ``w:shd`` element in the XML.

    Checks ``w:pPr`` (paragraph properties), ``w:tcPr`` (table-cell
    properties), and ``w:rPr`` (run properties) — whichever is present
    on the given element.

    This goes directly to the authoritative XML rather than using
    ``python-docx``'s high-level shading API, because the high-level
    API can return ``None`` both for "genuinely no shading" and "shading
    exists but the API failed to read it" — an ambiguity we cannot
    resolve without checking the XML anyway.
    """
    if element is None:
        return None

    for prop_tag in ("pPr", "tcPr", "rPr"):
        # ElementTree element doesn't expose 'find' in stubs correctly.
        props = element.find(f"{_WML_NS}{prop_tag}")  # type: ignore[attr-defined]
        if props is not None:
            shd = props.find(f"{_WML_NS}shd")
            if shd is not None:
                fill = shd.get(f"{_WML_NS}fill")
                if fill and fill.lower() not in ("auto", ""):
                    return fill
    return None


def _get_effective_bg(
    para_element: object,
    cell_element: object | None = None,
) -> tuple[int, int, int]:
    """Determine the effective background colour for text comparison.

    Checks paragraph shading first, then cell shading (if in a table),
    falling back to white ``(255, 255, 255)`` — the default document
    background.
    """
    # Paragraph-level shading
    fill_hex = _get_shading_fill_from_xml(para_element)
    if fill_hex:
        parsed = _parse_hex_colour(fill_hex)
        if parsed is not None:
            return parsed

    # Cell-level shading (if paragraph is inside a table cell)
    if cell_element is not None:
        fill_hex = _get_shading_fill_from_xml(cell_element)
        if fill_hex:
            parsed = _parse_hex_colour(fill_hex)
            if parsed is not None:
                return parsed

    return (255, 255, 255)


# ---------------------------------------------------------------------------
# DocxAdapter
# ---------------------------------------------------------------------------

class DocxAdapter(SourceAdapter):
    """Extract visible text and detect hidden content in Word documents.

    **Requires:** ``python-docx``.
    Install via ``pip install okf-guard[docx]``.
    """

    @property
    def format_name(self) -> str:
        """Return ``'docx'``."""
        return "docx"

    def extract(self, source: str | bytes) -> ExtractedContent:
        """Extract content from a ``.docx`` file.

        Args:
            source: A file path (``str``) or raw ``.docx`` bytes.

        Returns:
            ``ExtractedContent`` with visible text separated from
            hidden runs, plus metadata.

        Raises:
            FileNotFoundError: If *source* is a path that doesn't exist.
            ValueError: If the content cannot be parsed as a valid
                ``.docx`` file.
        """
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError(
                "DocxAdapter requires 'python-docx'.  "
                "Install it with:  pip install okf-guard[docx]"
            ) from exc

        path: str | None = None

        try:
            if isinstance(source, bytes):
                doc = Document(io.BytesIO(source))
            elif isinstance(source, str):
                if not os.path.isfile(source):
                    raise FileNotFoundError(
                        f"DOCX file not found: {source!r}"
                    )
                path = source
                doc = Document(source)
            else:
                raise TypeError(
                    f"DocxAdapter.extract() expects str or bytes, "
                    f"got {type(source).__name__}"
                )
        except (FileNotFoundError, TypeError, ImportError):
            raise
        except Exception as exc:
            raise ValueError(
                f"Failed to parse DOCX: {exc}"
            ) from exc

        visible_parts: list[str] = []
        hidden_spans: list[str] = []

        # Iterate all paragraphs in the main document body.
        for para_idx, paragraph in enumerate(doc.paragraphs, start=1):
            para_elem = paragraph._element
            bg = _get_effective_bg(para_elem)

            for run_idx, run in enumerate(paragraph.runs, start=1):
                text = run.text
                if not text.strip():
                    # Preserve whitespace-only runs in visible text
                    # for spacing, but don't bother checking for hidden.
                    visible_parts.append(text)
                    continue

                location = f"paragraph {para_idx}, run {run_idx}"

                # Check 1: Explicit hidden attribute — highest confidence
                # signal.  Must be explicitly True, not None (which
                # means "inherited / not set").
                if run.font.hidden is True:
                    hidden_spans.append(
                        f"[{location} — font.hidden] {text}"
                    )
                    continue

                # Check 2: Text colour matches effective background
                run_colour = _get_run_colour_rgb(run)
                if run_colour is not None and _colours_close(run_colour, bg):
                    hidden_spans.append(
                        f"[{location} — text color matches background] "
                        f"{text}"
                    )
                    continue

                # Not hidden — include in visible text.
                visible_parts.append(text)

        # Also iterate table cells for hidden content.
        # NOTE: Table content is checked for hidden text but also
        # included in visible text when not hidden, since tables are a
        # normal part of document content.
        for tbl_idx, table in enumerate(doc.tables, start=1):
            for row_idx, row in enumerate(table.rows, start=1):
                for cell_idx, cell in enumerate(row.cells, start=1):
                    cell_elem = cell._element
                    for para_idx_c, paragraph in enumerate(
                        cell.paragraphs, start=1,
                    ):
                        para_elem = paragraph._element
                        bg = _get_effective_bg(para_elem, cell_elem)

                        for run_idx, run in enumerate(
                            paragraph.runs, start=1,
                        ):
                            text = run.text
                            if not text.strip():
                                visible_parts.append(text)
                                continue

                            location = (
                                f"table {tbl_idx}, row {row_idx}, "
                                f"cell {cell_idx}, paragraph "
                                f"{para_idx_c}, run {run_idx}"
                            )

                            if run.font.hidden is True:
                                hidden_spans.append(
                                    f"[{location} — font.hidden] {text}"
                                )
                                continue

                            run_colour = _get_run_colour_rgb(run)
                            if (
                                run_colour is not None
                                and _colours_close(run_colour, bg)
                            ):
                                hidden_spans.append(
                                    f"[{location} — text color matches "
                                    f"background] {text}"
                                )
                                continue

                            visible_parts.append(text)

        metadata: dict[str, str] = {
            "format": self.format_name,
            "extracted_at": _utc_now_iso(),
        }
        if path is not None:
            metadata["path"] = path

        return ExtractedContent(
            text="".join(visible_parts),
            hidden_spans=hidden_spans,
            source_metadata=metadata,
        )
